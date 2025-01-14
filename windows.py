import sqlite3
import sys
from PyQt6 import uic, QtCore
from PyQt6.QtSql import QSqlDatabase
from PyQt6.QtWidgets import QMessageBox, QDialog, QTableWidgetItem, QFileDialog

import config
import helpers
import docx
import re
import datetime
# from main import main_window

planfact = []
socr_naming_list = []
str_financial_order = True

class Window:
    def __init__(self, ui) -> None:
        # get all table names in self.tables
        Form, Window = uic.loadUiType(ui)
        self.window = Window()
        self.form = Form()
        self.last_query = ""
        self.form.setupUi(self.window)
        self.cache = {"FOs": [''], "regions": [''], "cities": [''], "universities": ['']}
        # window.show()

    def connect_db(self):
        db = QSqlDatabase.addDatabase('QSQLITE')
        db.setDatabaseName(config.DB_NAME)
        if not db.open():
            print('Не удалось подключиться к базе')
            return False
        return db

    def get_data(self, table=None, query="SELECT * \nFROM {}", log=True):
        conn = sqlite3.connect(config.DB_NAME)
        cursor = conn.cursor()
        query = query.format(table) or query
        cursor.execute(query)
        data = cursor.fetchall()
        conn.close()
        if log and "Tp_nir" in query:
            self.last_query = query
        return data

    @staticmethod
    @helpers.send_args_inside_func
    def close(*windows):
        for w in windows:
            w.window.close()


class MainWindow(Window):
    def __init__(self, ui) -> None:
        if not self.connect_db():
            sys.exit(-1)

        tables = self.get_data(query="SELECT * FROM sqlite_master WHERE type='table'")
        self.tables = [t[1] for t in tables]
        super().__init__(ui)

    def show_table(self, table, title='', headers=None, column_widths=None, data=None):
        # check if table exists in database
        self.form.ViewWidget.show()
        if table != "Tp_nir":
            sort_toggle = False
            self.form.ViewWidget.setSortingEnabled(sort_toggle)
        if not data:
            data = self.get_data(table)
        self.form.ViewWidget.setRowCount(len(data))
        self.form.ViewWidget.setColumnCount(len(data[0]))
        for i in range(len(data)):
            for j in range(len(data[0])):
                item = QTableWidgetItem(str(data[i][j]))
                item.setTextAlignment(QtCore.Qt.AlignmentFlag.AlignLeft)
                self.form.ViewWidget.setItem(i, j, item)
        if headers:
            self.form.ViewWidget.setHorizontalHeaderLabels(headers)
        if column_widths and len(column_widths) == len(data[0]):
            for column, width in enumerate(column_widths):
                self.form.ViewWidget.setColumnWidth(column, width)

        self.form.ViewWidget.setSortingEnabled(False)
        self.form.label.setText(title)

    def sort_selected(self):
        item = self.form.comboBoxSort.currentText()
        query = self.last_query
        if not query or "Tp_nir" not in query:
            query = "SELECT * \nFROM Tp_nir"

        if item == "Сортировка по потенциальному ключу":
            sort_toggle = False
            self.form.ViewWidget.setSortingEnabled(sort_toggle)
            if "GROUP BY" not in query:
                query = "\n".join([query, "GROUP BY Tp_nir.codvuz, rnw"])
            else:
                query = "\n".join(query.split('\n')[:-1] + ["GROUP BY Tp_nir.codvuz, rnw"])
        else:
            if "GROUP BY" in query:
                query = "\n".join(query.split('\n')[:-1])

            sort_toggle = False if item == "Без сортировки" else True

        self.show_table('Tp_nir', 'Информация о НИР', headers=config.TP_NIR_HEADERS,
                        column_widths=config.TP_NIR_COLUMN_WIDTH,
                        data=self.get_data(query=query, log=False))
        self.form.ViewWidget.setSortingEnabled(sort_toggle)

    @helpers.send_args_inside_func
    def reset_filters(self, window):
        self.show_table('Tp_nir', 'Информация о НИР', config.TP_NIR_HEADERS, config.TP_NIR_COLUMN_WIDTH)
        self.form.comboBoxSort.setCurrentIndex(0)
        self.form.resetFiltersButton.setEnabled(False)
        window.condition = [None for _ in window.condition]
        window.reset_form()

    @helpers.send_args_inside_func
    def open_window(self, main_window, add_window, flag: bool):
        self.new_window = QDialog()
        self.ui_window = add_window
        self.ui_window.form.setupUi(self.new_window)
        self.new_window.show()
        if flag:
            self.ui_window.form.agreeButton.clicked.connect(self.add_button(main_window))
        else:
            selected_row = main_window.form.ViewWidget.selectionModel().selectedRows()
            selected_index = [index.row() for index in selected_row]

            if 0 < len(selected_index) < 2:

                data_stack = helpers.get_rows_from_table(selected_index[0], main_window.form.ViewWidget.model())
                if len(data_stack[4]) > 8:
                    cod_grnti = data_stack[4][:9]
                    cod_grnti_2 = data_stack[4][8:]
                    self.ui_window.form.cod_grnti.setText(cod_grnti)
                    self.ui_window.form.cod_grnti_2.setText(cod_grnti_2)

                self.ui_window.form.socr_naming.setEnabled(False)
                self.ui_window.form.socr_naming.setCurrentText(data_stack[3])

                self.ui_window.form.reg_number_nir.setDisabled(True)
                self.ui_window.form.reg_number_nir.setText(data_stack[1])

                dict_items = config.dict_character_nir.items()
                for key, value in dict_items:
                    if value == data_stack[2]:
                        character_nir = key

                self.ui_window.form.character_nir.setCurrentText(character_nir)

                self.ui_window.form.cod_grnti.setText(data_stack[4])
                self.ui_window.form.ruk_nir.setText(data_stack[5])
                self.ui_window.form.post.setText(data_stack[6])
                self.ui_window.form.financial.setValue(int(data_stack[7]))
                self.ui_window.form.naming_nir.setText(data_stack[8])

                self.ui_window.form.agreeButton.clicked.connect(
                    self.edit_button(int(data_stack[7]), selected_index[0], main_window))
            else:
                message_text = 'Не выбрана строка для редактирования или выбрано более одной!'
                self.ui_window.form.message = QMessageBox(QMessageBox.Icon.Critical, 'Ошибка', message_text)
                self.ui_window.form.message.show()
                self.new_window.close()

        self.ui_window.form.cancelButton.clicked.connect(self.new_window.close)

    @helpers.send_args_inside_func
    def add_button(self, main_window):
        """Функция добавления строки."""

        reg_number_nir = self.ui_window.form.reg_number_nir.text()
        character_nir = self.ui_window.form.character_nir.currentText()
        socr_naming = self.ui_window.form.socr_naming.currentText()
        cod_grnti = self.ui_window.form.cod_grnti.text()
        cod_grnti_2 = self.ui_window.form.cod_grnti_2.text()
        ruk_nir = self.ui_window.form.ruk_nir.text()
        post = self.ui_window.form.post.text()
        financial = self.ui_window.form.financial.value()
        naming_nir = self.ui_window.form.naming_nir.toPlainText()

        error = 0

        if (reg_number_nir != '' and
                character_nir != '' and
                socr_naming != '' and
                cod_grnti != '..' and
                ruk_nir != ' ..' and
                post != '' and
                naming_nir != ''):

            dict_items = config.dict_cod_vuz_socr_naming.items()
            for key, value in dict_items:
                if value == socr_naming:
                    cod_vuz = key
            data = self.get_data(query=f"""SELECT codvuz, rnw FROM Tp_nir""")

            set_cod_vuz_reg_number_nir = (cod_vuz, reg_number_nir)
            if set_cod_vuz_reg_number_nir in data:
                message_text = 'Связь код ВУЗа и рег.номера НИР не уникальна!'
                self.ui_window.form.message = QMessageBox(QMessageBox.Icon.Critical, 'Ошибка', message_text)
                self.ui_window.form.message.show()
                error = 1

            if financial <= 0:
                message_text = 'Значение планового финансирования не может быть меньше или равно нулю!'
                self.ui_window.form.message = QMessageBox(QMessageBox.Icon.Critical, 'Ошибка', message_text)
                self.ui_window.form.message.show()
                error = 1

            if ((len(cod_grnti) < 6 or ' ' in cod_grnti or len(cod_grnti) == 7) or
                    (cod_grnti_2 != '..' and (len(cod_grnti_2) < 6 or ' ' in cod_grnti_2 or len(cod_grnti_2) == 7))):
                message_text = f"""Проверьте правильность написания кода ГРНТИ:\n
                            1) Не должно быть пробелов в коде\n 
                            2) Минимальный код состоит из четырех цифр\n
                            3) Код не может состоять из пяти цифр
                            """
                self.ui_window.form.message = QMessageBox(QMessageBox.Icon.Critical, 'Ошибка', message_text)
                self.ui_window.form.message.show()
                error = 1

            if config.dict_cod_vuz_socr_naming[cod_vuz] != socr_naming:
                message_text = f'Данному коду ВУЗа соответсвует название: {config.dict_cod_vuz_socr_naming[cod_vuz]}'
                self.ui_window.form.message = QMessageBox(QMessageBox.Icon.Critical, 'Ошибка', message_text)
                self.ui_window.form.message.show()
                error = 1

            if error == 0:
                if len(cod_grnti) == 6:
                    cod_grnti = cod_grnti[:5]

                if cod_grnti_2 != '..':
                    if len(cod_grnti_2) == 6:
                        cod_grnti_2 = cod_grnti_2[:5]
                    cod_grnti = cod_grnti + ';' + cod_grnti_2

                query_tp_nir = f"""INSERT INTO Tp_nir (codvuz, rnw, f1, z2, f10, f6, f7, f18, f2) 
                    VALUES ({cod_vuz},'{reg_number_nir}','{config.dict_character_nir[character_nir]}','{socr_naming}','{cod_grnti}', '{ruk_nir}', '{post}',{financial},'{naming_nir}');"""
                index = helpers.query_change_db(query_tp_nir)

                query_tp_fv_set_codvuz = "SELECT codvuz FROM Tp_fv;"
                set_codvuz = helpers.get_headers(table='Tp_fv',
                                                 query=query_tp_fv_set_codvuz)
                set_codvuz = [i[0] for i in set_codvuz]

                if cod_vuz not in set_codvuz:
                    query_tp_fv = f"""
                                INSERT INTO Tp_fv (codvuz, z2, z3, z18, numworks)
                                VALUES ({cod_vuz},'{socr_naming}', {financial}, 'None', {1});
                                """
                    helpers.query_change_db(query_tp_fv)
                else:
                    query_tp_fv = f"""
                                UPDATE Tp_fv
                                SET z3 = z3 + {financial}, numworks = numworks + 1
                                WHERE codvuz = '{cod_vuz}';
                                """
                    helpers.query_change_db(query_tp_fv)
                main_window.show_table('Tp_nir', 'Информация о НИР', config.TP_NIR_HEADERS, config.TP_NIR_COLUMN_WIDTH)

                self.new_window.close()
                main_window.form.ViewWidget.selectRow(index - 1)
        else:
            message_text = 'Заполните все поля!'
            self.ui_window.form.message = QMessageBox(QMessageBox.Icon.Critical, 'Ошибка заполнения полей',
                                                      message_text)
            self.ui_window.form.message.show()

    @helpers.send_args_inside_func
    def edit_button(self, fin0, index, main_window):
        """Функция редактирования строки."""

        socr_naming = self.ui_window.form.socr_naming.currentText()
        reg_number_nir = self.ui_window.form.reg_number_nir.text()
        character_nir = self.ui_window.form.character_nir.currentText()
        cod_grnti = self.ui_window.form.cod_grnti.text()
        cod_grnti_2 = self.ui_window.form.cod_grnti_2.text()
        ruk_nir = self.ui_window.form.ruk_nir.text()
        post = self.ui_window.form.post.text()
        financial = self.ui_window.form.financial.value()
        naming_nir = self.ui_window.form.naming_nir.toPlainText()

        error = 0
        if (
                character_nir != '' and
                cod_grnti != '..' and
                ruk_nir != ' ..' and
                post != '' and
                naming_nir != ''):

            if financial <= 0:
                message_text = 'Значение планового финансирования не может быть меньше или равно нулю!'
                self.ui_window.form.message = QMessageBox(QMessageBox.Icon.Critical, 'Ошибка', message_text)
                self.ui_window.form.message.show()
                error = 1

            if ((len(cod_grnti) < 6 or ' ' in cod_grnti or len(cod_grnti) == 7) or
                    (cod_grnti_2 != '..' and (len(cod_grnti_2) < 6 or ' ' in cod_grnti_2 or len(cod_grnti_2) == 7))):
                message_text = f"""Проверьте правильность написания кода ГРНТИ:\n
                            1) Не должно быть пробелов в коде\n 
                            2) Минимальный код состоит из четырех цифр\n
                            3) Код не может состоять из пяти цифр
                            """
                self.ui_window.form.message = QMessageBox(QMessageBox.Icon.Critical, 'Ошибка', message_text)
                self.ui_window.form.message.show()
                error = 1

            if error == 0:

                if len(cod_grnti) == 6:
                    cod_grnti = cod_grnti[:5]

                if cod_grnti_2 != '..':
                    if len(cod_grnti_2) == 6:
                        cod_grnti_2 = cod_grnti_2[:5]
                    cod_grnti = cod_grnti + ';' + cod_grnti_2
                dict_items = config.dict_cod_vuz_socr_naming.items()
                for key, value in dict_items:
                    if value == socr_naming:
                        cod_vuz = key
                query_tp_nir = f"""UPDATE Tp_nir 
                                SET f1 = '{config.dict_character_nir[character_nir]}', 
                                    f10 = '{cod_grnti}', 
                                    f6 = '{ruk_nir}',
                                    f7 = '{post}',
                                    f18 = {financial},
                                    f2 =  '{naming_nir}'
                                WHERE codvuz = '{cod_vuz}' AND rnw = '{reg_number_nir}';"""
                helpers.query_change_db(query_tp_nir)

                if fin0 != financial:
                    query_tp_fv = f"""
                                UPDATE Tp_fv
                                SET z3 = z3 - {fin0} + {financial}
                                WHERE codvuz = '{cod_vuz}';
                                """
                    helpers.query_change_db(query_tp_fv)
                main_window.show_table('Tp_nir', 'Информация о НИР', config.TP_NIR_HEADERS, config.TP_NIR_COLUMN_WIDTH)

                self.new_window.close()
                main_window.form.ViewWidget.selectRow(index)
        else:
            message_text = 'Заполните все поля!'
            self.ui_window.form.message = QMessageBox(QMessageBox.Icon.Critical, 'Ошибка заполнения полей',
                                                      message_text)
            self.ui_window.form.message.show()

    @helpers.send_args_inside_func
    def delete_button(self, main_window):
        """Функция удаления строки."""

        selected_rows = main_window.form.ViewWidget.selectionModel().selectedRows()
        selected_indexes = [index.row() for index in selected_rows]
        main_window.form.message = QMessageBox()
        if len(selected_indexes) > 0:
            main_window.form.message.setText("Вы уверены, что хотите удалить выделенные записи?")
            main_window.form.message.setStandardButtons(QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
            if main_window.form.message.exec() == main_window.form.message.StandardButton.Yes:
                stack = []
                for index in selected_indexes:
                    stack.append(helpers.get_rows_from_table(index, main_window.form.ViewWidget.model()))

                for i in range(len(stack)):
                    query_tp_nir = f"DELETE FROM Tp_nir WHERE codvuz = '{stack[i][0]}' AND rnw = '{stack[i][1]}';"
                    helpers.query_change_db(query=query_tp_nir)

                    query_tp_fv_get_numworks = f"SELECT numworks FROM Tp_fv WHERE codvuz = '{stack[i][0]}';"
                    numworks = helpers.get_headers('Tp_fv', query_tp_fv_get_numworks)
                    numworks = [i[0] for i in numworks]

                    if numworks[0] == 1:
                        query_tp_fv = f"""
                                        DELETE FROM Tp_fv
                                        WHERE codvuz = '{stack[i][0]}';
                                        """
                        helpers.query_change_db(query=query_tp_fv)

                    else:
                        query_tp_fv = f"""
                                        UPDATE Tp_fv
                                        SET z3 = z3 - {stack[i][7]}, numworks = numworks - 1
                                        WHERE codvuz = '{stack[i][0]}';
                                        """
                        helpers.query_change_db(query=query_tp_fv)

                    main_window.form.ViewWidget.model().removeRow(index)
                    main_window.show_table('Tp_nir', 'Информация о НИР', config.TP_NIR_HEADERS,
                                           config.TP_NIR_COLUMN_WIDTH)

            else:
                message_text = 'Удаление отменено'
                main_window.form.message = QMessageBox(QMessageBox.Icon.Critical, 'Ошибка удаления', message_text)
                main_window.form.message.show()
        else:
            message_text = 'Не выбрана запись для удаления'
            main_window.form.message = QMessageBox(QMessageBox.Icon.Critical, 'Ошибка удаления', message_text)
            main_window.form.message.show()

    @helpers.send_args_inside_func
    def showFrame(self, qmenu):
        if qmenu == 'Data':
            self.form.analFrame.hide()
            self.form.FinancialFrame.hide()
            self.form.helpframe.hide()
            self.form.Frame.show()
        elif qmenu == 'Analyses':
            self.form.Frame.hide()
            self.form.FinancialFrame.hide()
            self.form.helpframe.hide()
            self.form.analFrame.show()
        elif qmenu == 'Financial':
            self.form.Frame.hide()
            self.form.analFrame.hide()
            self.form.helpframe.hide()
            self.form.FinancialFrame.show()
        elif qmenu == 'Help':
            self.form.Frame.hide()
            self.form.FinancialFrame.hide()
            self.form.analFrame.hide()
            self.form.helpframe.show()


    def constructor_table(self, tablename, column_widths, column_headers, data, total=True):
        ###{tablenme = 'VUZ'/'har'/'grnti'
        # data = [column1, column2, ..., columnN] of 1-3 tables
        # itogo = True if need to add itogo to table}###
        getattr(self.form, f"{tablename}").clear()
        getattr(self.form, f"{tablename}").show()
        if total:
            getattr(self.form, f"{tablename}").setRowCount(len(data[0]) + 1)
        else:
            getattr(self.form, f"{tablename}").setRowCount(len(data[0]))
        getattr(self.form, f"{tablename}").setColumnCount(len(column_headers))
        itogo_1, itogo_2 = 0, 0
        for column in range(len(data)):
            for row in range(len(data[column])):
                item = QTableWidgetItem(str(data[column][row]))
                if type(data[column][row]) == int:
                    item.setTextAlignment(QtCore.Qt.AlignmentFlag.AlignRight)
                else:
                    item.setTextAlignment(QtCore.Qt.AlignmentFlag.AlignLeft)
                getattr(self.form, f"{tablename}").setItem(row, column, item)
                if column == 1 and total:
                    itogo_1 += data[column][row]
                elif column == 2 and total:
                    itogo_2 += data[column][row]
            getattr(self.form, f"{tablename}").setColumnWidth(column, column_widths[column])
        getattr(self.form, f"{tablename}").setHorizontalHeaderLabels(column_headers)
        getattr(self.form, f"{tablename}").sortItems(0, QtCore.Qt.SortOrder.AscendingOrder)
        if total:
            itogo_t = QTableWidgetItem('Итого:')
            itogo_t.setTextAlignment(QtCore.Qt.AlignmentFlag.AlignLeft)
            itogo_1_t = QTableWidgetItem(str(itogo_1))
            itogo_1_t.setTextAlignment(QtCore.Qt.AlignmentFlag.AlignRight)
            getattr(self.form, f"{tablename}").setItem(row + 1, 0, itogo_t)
            getattr(self.form, f"{tablename}").setItem(row + 1, 1, itogo_1_t)
            if itogo_2:
                itogo_2_t = QTableWidgetItem(str(itogo_2))
                itogo_2_t.setTextAlignment(QtCore.Qt.AlignmentFlag.AlignRight)
                getattr(self.form, f"{tablename}").setItem(row + 1, 2, itogo_2_t)


    @helpers.send_args_inside_func
    def create_tables(self, namewidget):
        qtable = getattr(self.form, f"{namewidget}")
        data, items_z2, har, code, code_list, numworks_1tab, sumfin_1tab, harac = [], [], [], [], [], [], [], []
        numworks_2tab, sumfin_2tab, numworks_3tab, sumfin_3tab, namegrnti = [], [], [], [], []
        for row in range(qtable.rowCount()):
            dat = []
            for col in range(qtable.columnCount()):
                item = qtable.item(row, col)
                dat.append(item.text())
            data.append(dat)
        for k in range(len(data)):
            items_z2.append(data[k][3])
            har.append(data[k][2])
            code.append(data[k][4])
        for i in range(len(code)):
            if len(code[i]) <= 10:
                code1 = code[i][0:2]
            else:
                code1 = code[i][0:2]
                code2 = code[i][9:11]
                code_list.append(code2)
            code_list.append(code1)
        code_list = list(set(code_list))
        items_z2 = list(set(items_z2))
        har = list(set(har))
        for i in range(len(items_z2)):
            numworks_1 = 0
            sumfin_1 = 0
            for j in range(len(data)):
                if data[j][3] == items_z2[i]:
                    numworks_1 += 1
                    sumfin_1 += int(data[j][7])
            numworks_1tab.append(numworks_1)
            sumfin_1tab.append(sumfin_1)
        table1 = [items_z2, numworks_1tab, sumfin_1tab]
        for i in range(len(har)):
            if har[i] == 'Ф':
                harac.append('Фундаментальное исследование')
            elif har[i] == 'П':
                harac.append('Прикладное исследование')
            else:
                harac.append('Экспериментальная разработка')
            numworks_2 = 0
            sumfin_2 = 0
            for j in range(len(data)):
                if data[j][2] == har[i]:
                    numworks_2 += 1
                    sumfin_2 += int(data[j][7])
            numworks_2tab.append(numworks_2)
            sumfin_2tab.append(sumfin_2)
        table2 = [harac, numworks_2tab, sumfin_2tab]
        for j in range(len(code_list)):
            numworks_3 = 0
            sumfin_3 = 0
            for i in range(len(data)):
                s = data[i][4]
                match = re.match(r"(\d{2}).*", s)
                if (code_list[j] == match.group(1)):
                    numworks_3 += 1
                    sumfin_3 += int(data[i][7])
                if (len(data[i][4]) >= 12):
                    match2 = re.match(r".{9}(\d{2})", s)
                    if (code_list[j] == match2.group(1)) and (code_list[j] != match.group(1)):
                        numworks_3 += 1
                        sumfin_3 += int(data[i][7])
            numworks_3tab.append(numworks_3)
            sumfin_3tab.append(sumfin_3)
            namegrnti.append(self.get_data('database.db',
                                           query=f"""SELECT rubrika FROM grntirub WHERE codrub == '{code_list[j]}'""")[0][0])
        
        table3 = [code_list, namegrnti, numworks_3tab, sumfin_3tab]
        self.constructor_table('NirbVUZ', config.COLUMN_WIDTHS_NirbVUZ, config.HEADERS_NirbVUZ, table1, True)
        self.constructor_table('Nirbhar', config.COLUMN_WIDTHS_Nirbhar, config.HEADERS_Nirbhar, table2, True)
        self.constructor_table('Nirbgrnti', config.COLUMN_WIDTHS_Nirbgrnti, config.HEADERS_Nirbgrnti, table3, False)

    @helpers.send_args_inside_func
    def save_to_docx(self, qtable1, window, utverdit=None, main_window=None):
        flag = True

        if utverdit != None:
            if (planfact == [] and socr_naming_list == []) or str_financial_order is True:
                message_text = 'Не заполнены поля для рассчета фактического плана финансирования!'
                main_window.form.message = QMessageBox(QMessageBox.Icon.Critical, 'Ошибка заполнения полей',
                                                       message_text)
                main_window.form.message.show()
                flag = False

        if flag:
            cond = window.condition
            filename = QFileDialog.getSaveFileName(None, 'Save File', '.', 'Документ Microsoft Word (*.docx)')[0]
            if filename == '':
                return
            if qtable1 == 'Financial':
                qtable = getattr(self.form, f"{qtable1}Utverdit")
            else:
                qtable = getattr(self.form, f"Nirb{qtable1}")
            doc = docx.Document()
            fil_name = ['Федеральный округ: ', 'Субъект федерации: ', 'Город: ', 'ВУЗ: ', 'Первые цифры кода ГРНТИ: ']
            if qtable1 == 'VUZ':
                doc.add_heading('Распределение НИР по ВУЗам')
            elif qtable1 == 'grnti':
                doc.add_heading('Распределение НИР по коду ГРНТИ')
            elif qtable1 == 'har':
                doc.add_heading('Распределение НИР по характеру')
            elif qtable1 == 'Financial':
                doc.add_heading('Распоряжение по финансированию ВУЗов')
                cond = 0
            if cond:
                doc.add_paragraph('Примененные фильтры')
                for fp in range(len(cond)):
                    if cond[fp]:
                        doc.add_paragraph(fil_name[fp] + cond[fp])
            table = doc.add_table(rows=qtable.rowCount() + 1, cols=qtable.columnCount())
            table.style = 'Table Grid'
            for i in range(qtable.columnCount()):
                table.cell(0, i).text = qtable.horizontalHeaderItem(i).text()
            for row in range(qtable.rowCount()):
                for col in range(qtable.columnCount()):
                    item = qtable.item(row, col)
                    table.cell(row + 1, col).text = item.text() if item else ""
            if qtable1 != 'Financial':
                doc.add_paragraph('Общее число НИР: ' + str(
                    self.get_data('database.db', f"""SELECT COUNT(f18) FROM Tp_nir""", log=False)[0][0]))
                doc.add_paragraph('Общая сумма планового финансирования: ' + str(
                    self.get_data('database.db', f"""SELECT SUM(f18) FROM Tp_nir""", log=False)[0][0]))
            else:
                doc.add_paragraph('Распоряжение от: '+ datetime.date.today().isoformat())
            doc.save(f"{filename}")

    @helpers.send_args_inside_func
    def print_filter(self, window):
        cond = window.condition
        fil_name = ['Федеральный округ: ', 'Субъект федерации: ', 'Город: ', 'ВУЗ: ', 'Первые цифры кода ГРНТИ: ']
        self.form.ApplFilter.clear()
        if cond:
            for fp in range(len(cond)):
                if cond[fp]:
                    self.form.ApplFilter.append(fil_name[fp] + cond[fp])
        summ = self.get_data('database.db', query=f"""SELECT SUM(f18) FROM Tp_nir""", log=False)
        count = self.get_data('database.db', query=f"""SELECT COUNT(f18) FROM Tp_nir""", log=False)
        self.form.nirsum.setText('Общее число НИР: ' + str(count[0][0]))
        self.form.finsum.setText('Общая сумма финансирования: ' + str(summ[0][0]))

    @helpers.send_args_inside_func
    def create_finacial_tables(self, main_window):
        """Функция создания таблицы финансирования"""

        query_to_tp_fv = """SELECT z2 FROM Tp_fv"""
        socr_naming = self.get_data('database.db', query=query_to_tp_fv, log=False)
        socr_naming_list = []
        for i in range(len(socr_naming)):
            socr_naming_list.append(socr_naming[i][0])

        query_to_tp_fv = """SELECT COUNT(*) FROM Tp_fv"""
        sumfact = self.get_data('database.db', query=query_to_tp_fv, log=False)
        sumfact_list = [0]*sumfact[0][0]
        
        data = [socr_naming_list, sumfact_list]

        self.constructor_table('FinancialUtverdit', (180, 160), ('Сокращенное название ВУЗа', 'Размер финансирования'), data)

    @helpers.send_args_inside_func
    def financing_order(self, main_window):
        """Функция, отвечающая за окно финансовое распоряжение."""
        
        # заполнение полей
        global str_financial_order
        str_financial_order = True

        target_nir = self.get_data('database.db', query=f"""SELECT SUM(numworks) FROM Tp_fv""", log=False)
        main_window.form.targetNir.setText(str(target_nir[0][0]))

        target_vuz = self.get_data('database.db', query=f"""SELECT COUNT(*) FROM Tp_fv""", log=False)
        main_window.form.targetVuz.setText(str(target_vuz[0][0]))

        targetsumplan = self.get_data('database.db', query=f"""SELECT SUM(z3) FROM Tp_fv""", log=False)
        main_window.form.targetsumplan.setText(str(targetsumplan[0][0]))

        targetsumfact = self.get_data('database.db', query=f"""SELECT SUM(z18) FROM Tp_fv""", log=False)
        targetsumfact = targetsumfact[0][0]
        if targetsumfact == None:
            targetsumfact = 0
        main_window.form.targetsumfact.setText(str(targetsumfact))

        try: 
            procfact = round(targetsumfact / targetsumplan[0][0] * 100, 2)
        except ZeroDivisionError:
            procfact = 0
        if procfact == -0.0:
            procfact = 0.0
        main_window.form.procfact.setText(str(procfact))
        
        # заполнение полей 'сумма к распределению' или '% от планового'  
        main_window.form.Calculation.clicked.connect(main_window.calculation(main_window,
                                                                             targetsumplan[0][0]))
    
    @helpers.send_args_inside_func
    def calculation(self, main_window, targetsumplan: int):
        """Функция кнопки 'Рассчитать'. """

        sumplan = main_window.form.sumplan.text()
        procplan = main_window.form.procplan.text()

        flag = 0
        # глобальная переменная, которая обозначает, что пользователь зашел на страницу финансирования
        # и больше никаких действий не сделал
        global str_financial_order
        str_financial_order =False

        if sumplan != '' and procplan != '':
            sumplan = round(targetsumplan*float(procplan)//100)
            main_window.form.sumplan.setText(str(sumplan))

        if sumplan == '' and procplan == '':
            message_text = 'Заполните поля!'
            main_window.form.message = QMessageBox(QMessageBox.Icon.Critical, 'Ошибка заполнения полей',
                                                   message_text)
            main_window.form.message.show()

            flag = 1 

        if procplan != '' and sumplan == '':
            print(sumplan)
            if round(float(procplan),2) < 0.01 and round(float(procplan),2) > -0.01:
                message_text = 'Слишком маленький процент финансирования / обратного финансирования (меньше 0.01 %)!'
                main_window.form.message = QMessageBox(QMessageBox.Icon.Critical, 'Ошибка заполнения полей',
                                                       message_text)
                main_window.form.message.show()
                main_window.form.sumplan.setText('')

                flag = 1

        if procplan == '' and sumplan != '':
            procplan = round(int(sumplan) / targetsumplan*100, 2)
            main_window.form.procplan.setText(str(procplan))
            if round(float(procplan),2) < 0.01 and round(float(procplan),2) > -0.01:
                message_text = 'Слишком маленький процент финансирования / обратного финансирования (меньше 0.01 %)!'
                main_window.form.message = QMessageBox(QMessageBox.Icon.Critical, 'Ошибка заполнения полей',
                                                       message_text)
                main_window.form.message.show()
                main_window.form.procplan.setText('')

                flag = 1

        if sumplan == '' and procplan != '' and flag == 0:
            print(sumplan)
            sumplan = round(targetsumplan*float(procplan)//100)
            main_window.form.sumplan.setText(str(sumplan))

        if flag == 0:
            targetsumplan = self.get_data('database.db', query=f"""SELECT SUM(z3) FROM Tp_fv""", log=False)
            targetsumplan = targetsumplan[0][0]

            targetsumfact = self.get_data('database.db', query=f"""SELECT SUM(z18) FROM Tp_fv""", log=False)
            targetsumfact = targetsumfact[0][0]
            try: 
                procfact = round(targetsumfact / targetsumplan * 100, 2)
            except ZeroDivisionError:
                procfact = 0
            if procfact == -0.0:
                procfact = 0.0

            global main_percent_procfact 
            main_percent_procfact=procfact # float с округлением до 2х цифр после запятой
            
            query_to_tp_fv = """SELECT z2, z3 FROM Tp_fv"""
            data_tp_fv_z2_z3 = self.get_data('database.db', query=query_to_tp_fv, log=False)

            global socr_naming_list
            socr_naming_list = []

            global planfact
            planfact = []

            global dop_percent_procplan 
            dop_percent_procplan = round(float(procplan), 2)

            for i in range(len(data_tp_fv_z2_z3)):
                socr_naming_list.append(data_tp_fv_z2_z3[i][0])
                planfact.append(round(data_tp_fv_z2_z3[i][1]*float(procplan)//100))
            data = [socr_naming_list, planfact]
            self.constructor_table('FinancialUtverdit', (180, 160), ('Сокращенное название ВУЗа', 'Размер финансирования'), data)



    @helpers.send_args_inside_func
    def utverdit(self, main_window):
        """Кнопка утверждения приказа."""
        
        if (planfact == [] and socr_naming_list == []) or str_financial_order is True:
            message_text = 'Не заполнены поля для рассчета фактического плана финансирования!'
            main_window.form.message = QMessageBox(QMessageBox.Icon.Critical, 'Ошибка заполнения полей',
                                                   message_text)
            main_window.form.message.show()
        else:
            if (dop_percent_procplan + main_percent_procfact) < 0:
                message_text = 'Суммарный процент от планового финансирования не может быть меньше нуля!'
                main_window.form.message = QMessageBox(QMessageBox.Icon.Critical, 'Ошибка заполнения полей',
                                                        message_text)
                main_window.form.message.show()
            else: 

                helpers.query_change_db(fact=True, socr_naming=True)
                targetsumplan = self.get_data('database.db', query=f"""SELECT SUM(z3) FROM Tp_fv""", log=False)
                targetsumfact = self.get_data('database.db', query=f"""SELECT SUM(z18) FROM Tp_fv""", log=False)
                targetsumfact = targetsumfact[0][0]
                
                try: 
                    procfact = round(targetsumfact / targetsumplan[0][0] * 100, 2)
                except ZeroDivisionError:
                    procfact = 0
                if procfact == -0.0:
                    procfact = 0.0

                main_window.form.targetsumfact.setText(str(targetsumfact))
                main_window.form.procfact.setText(str(procfact))
                
                message_text = ("Приказ о финансировании успешно утвержден!")
                main_window.form.message = QMessageBox(QMessageBox.Icon.Information, 'Успешное выполение', message_text)
                main_window.form.message.setStandardButtons(QMessageBox.StandardButton.Ok)
                main_window.form.message.show()


class FilterWindow(Window):
    def __init__(self, ui) -> None:
        super().__init__(ui)
        self.boxes_data = {"FO": [""], "Region": [""], "City": [""], "University": [""]}
        self.condition = []
        self.column2widget = dict(zip(["region", "oblname", "city", "VUZ.z2"],
                                      self.boxes_data))
        self.query_template = '\n'.join(["SELECT DISTINCT {select_column}",
                                         "FROM Tp_nir JOIN VUZ ON Tp_nir.codvuz = VUZ.codvuz",
                                         "WHERE {column} = \"{value}\""])
        self.condition = []
        for data in self.get_data(query="\n".join([f"SELECT DISTINCT {', '.join(self.column2widget)}",
                                                   "FROM VUZ JOIN Tp_nir ON VUZ.codvuz = Tp_nir.codvuz"])):
            self.boxes_data['FO'].append(data[0])
            self.boxes_data['Region'].append(data[1])
            self.boxes_data['City'].append(data[2])
            self.boxes_data['University'].append(data[3])

        self.boxes_data = {k: sorted(set(v)) for k, v in self.boxes_data.items()}
        for level, data in self.boxes_data.items():
            getattr(self.form, f"comboBox{level}").addItems(data)

    def update_filter_data(self):
        self.condition = []
        boxes_data = {"FO": [""], "Region": [""], "City": [""], "University": [""]}
        for data in self.get_data(query="\n".join([f"SELECT DISTINCT {', '.join(self.column2widget)}",
                                                   "FROM VUZ JOIN Tp_nir ON VUZ.codvuz = Tp_nir.codvuz"])):
            boxes_data['FO'].append(data[0])
            boxes_data['Region'].append(data[1])
            boxes_data['City'].append(data[2])
            boxes_data['University'].append(data[3])

        boxes_data = {k: sorted(set(v)) for k, v in boxes_data.items()}
        levels_for_update = [level for level, data in boxes_data.items() if (
                    1 < len(self.boxes_data[level]) != len(data))]
        for key in levels_for_update:
            buffer = getattr(self.form, f"comboBox{key}").currentText()
            getattr(self.form, f"comboBox{key}").clear()
            getattr(self.form, f"comboBox{key}").addItems(boxes_data[key])
            i = getattr(self.form, f"comboBox{key}").findText(buffer)
            if i < 0:
                i = 0
            getattr(self.form, f"comboBox{key}").setCurrentIndex(i)
            self.boxes_data[key] = boxes_data[key]

        self.window.show()

    def get_combobox_values(self, target_column):
        query = []
        for column, filter_level in self.column2widget.items():
            if column == target_column:
                break
            value = getattr(self.form, f"comboBox{filter_level}").currentText()
            if value:
                query.append(self.query_template.format(column=column,
                                                        value=value,
                                                        select_column=target_column))

        if query:
            query = "\nINTERSECT\n".join(set(query))
        else:
            query = f"SELECT DISTINCT {target_column} \nFROM VUZ JOIN Tp_nir ON Tp_nir.codvuz = VUZ.codvuz"

        return sorted([element[0] for element in self.get_data(query=query)])

    @helpers.send_args_inside_func
    def combobox_filter(self, target):
        data = self.get_combobox_values(target)
        filter_level = self.column2widget[target]
        self.boxes_data[filter_level] = [""] + data
        getattr(self.form, f"comboBox{filter_level}").clear()
        getattr(self.form,
                f"comboBox{filter_level}").addItems(self.boxes_data[filter_level])

    @helpers.send_args_inside_func
    def apply_filter(self, window: MainWindow):
        self.condition = []
        query_template = "\n".join(["SELECT Tp_nir.*",
                                    "FROM Tp_nir JOIN VUZ ON Tp_nir.codvuz = VUZ.codvuz",
                                    "WHERE {column} {sign} \"{value}\""])
        grnti = self.form.lineEdit.text()
        queries = []
        if grnti.isdigit():
            queries.append("\nUNION\n".join(
                [query_template.format(column="f10", sign="LIKE", value=f"{grnti}%"),
                 query_template.format(column="f10", sign="LIKE", value=f"%;{grnti}%")]))
            self.condition.append(grnti)

        for column, box_name in sorted(self.column2widget.items()):
            value = getattr(self.form, f"comboBox{box_name}").currentText()
            self.condition.append(value or None)
            if value:
                queries.append(query_template.format(column=column, sign="=", value=value))

        self.condition = self.condition[::-1]
        if not queries:
            return
        data = window.get_data(query="\nINTERSECT\n".join(queries))
        window.show_table('Tp_nir', 'Информация о НИР', config.TP_NIR_HEADERS,
                          config.TP_NIR_COLUMN_WIDTH, data)
        window.form.resetFiltersButton.setEnabled(True)
        self.window.hide()

    def reset_form(self):
        self.form.lineEdit.clear()
        for box_name in list(self.boxes_data)[::-1]:
            if getattr(self.form, f"comboBox{box_name}").currentIndex():
                getattr(self.form, f"comboBox{box_name}").setCurrentIndex(0)
